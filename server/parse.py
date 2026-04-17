from __future__ import annotations

import os
import re
from pathlib import Path
from typing import TypedDict

from pypdf import PdfReader


class ParsedDocument(TypedDict):
	document_type: str
	file_name: str
	content: str


_EASYOCR_READER = None


def _use_gpu() -> bool:
	use_gpu_env = os.getenv("USE_GPU", "auto").strip().lower()
	if use_gpu_env in {"0", "false", "no", "off"}:
		return False
	if use_gpu_env in {"1", "true", "yes", "on"}:
		return True

	try:
		import torch

		return bool(torch.cuda.is_available())
	except Exception:
		return False


def _get_ocr_reader():
	global _EASYOCR_READER
	if _EASYOCR_READER is None:
		import easyocr

		_EASYOCR_READER = easyocr.Reader(["en"], gpu=_use_gpu())
	return _EASYOCR_READER


def _clean_extracted_text(text: str) -> str:
	"""
	Remove binary / non-printable content that leaks when DOCX or image-based PDFs
	are processed. Without this, raw DOCX XML bytes end up in the LLM prompt and
	cause the model to echo garbage or describe binary structure instead of writing
	a real BRD.
	"""
	if not text:
		return ""

	# Strip non-printable / binary-looking bytes (keep standard Unicode text ranges)
	cleaned = re.sub(
		r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uD7FF\uF900-\uFDCF\uFDF0-\uFFEF]+",
		" ",
		text,
	)
	# Collapse long runs of whitespace / repeated spaces
	cleaned = re.sub(r"[ \t]{4,}", "   ", cleaned)
	cleaned = re.sub(r"\n{4,}", "\n\n\n", cleaned)
	return cleaned.strip()


def _looks_like_binary(text: str) -> bool:
	"""
	Return True if the extracted text is mostly binary garbage —
	e.g. when pypdf reads a fully scanned/image PDF with no embedded text.
	"""
	if not text:
		return True
	printable = sum(1 for c in text if c.isprintable() or c in "\n\r\t")
	return (printable / max(len(text), 1)) < 0.60


def _extract_text_from_image(file_path: Path) -> str:
	reader = _get_ocr_reader()
	result = reader.readtext(str(file_path), detail=0, paragraph=True)
	raw = "\n".join([line.strip() for line in result if line and line.strip()])
	return _clean_extracted_text(raw)


def _extract_text_from_pdf(file_path: Path) -> str:
	# Try OCR-first for scanned PDFs when PyMuPDF is available.
	try:
		import fitz  # type: ignore

		reader = _get_ocr_reader()
		doc = fitz.open(str(file_path))
		ocr_lines: list[str] = []
		for page in doc:
			pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
			img_bytes = pix.tobytes("png")
			lines = reader.readtext(img_bytes, detail=0, paragraph=True)
			ocr_lines.extend([line.strip() for line in lines if line and line.strip()])
		doc.close()

		if ocr_lines:
			return _clean_extracted_text("\n".join(ocr_lines))
	except Exception:
		pass

	# Fallback to embedded text extraction.
	pdf_reader = PdfReader(str(file_path))
	text_lines: list[str] = []
	for page in pdf_reader.pages:
		text = (page.extract_text() or "").strip()
		if text:
			text_lines.append(text)

	raw = "\n\n".join(text_lines)

	# If the result looks binary (scanned PDF, no embedded text), return empty
	# so the caller knows there's nothing useful rather than feeding garbage to LLM.
	if _looks_like_binary(raw):
		return ""

	return _clean_extracted_text(raw)


def _extract_text_from_docx(file_path: Path) -> str:
	"""
	Extract readable text from a .docx file without leaking raw XML/binary bytes.
	Uses python-docx when available, otherwise falls back to zip+XML parsing.
	"""
	try:
		from docx import Document  # type: ignore

		doc = Document(str(file_path))
		paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
		return _clean_extracted_text("\n\n".join(paragraphs))
	except Exception:
		pass

	# Fallback: unzip and grab document.xml text nodes manually
	try:
		import zipfile
		import xml.etree.ElementTree as ET

		ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
		texts: list[str] = []
		with zipfile.ZipFile(str(file_path), "r") as z:
			if "word/document.xml" in z.namelist():
				xml_bytes = z.read("word/document.xml")
				root = ET.fromstring(xml_bytes)
				for t_elem in root.iter(f"{ns}t"):
					if t_elem.text and t_elem.text.strip():
						texts.append(t_elem.text.strip())
		return _clean_extracted_text(" ".join(texts))
	except Exception:
		return ""


def _extract_text_generic(file_path: Path) -> str:
	try:
		raw = file_path.read_text(encoding="utf-8")
	except UnicodeDecodeError:
		raw = file_path.read_text(encoding="latin-1", errors="ignore")
	return _clean_extracted_text(raw)


def parse_file_content(file_path: Path) -> str:
	suffix = file_path.suffix.lower()

	if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"}:
		return _extract_text_from_image(file_path)

	if suffix == ".pdf":
		return _extract_text_from_pdf(file_path)

	# Handle Word documents properly — never read raw bytes
	if suffix in {".docx", ".doc"}:
		return _extract_text_from_docx(file_path)

	if suffix in {".txt", ".md", ".csv", ".json", ".log"}:
		return _extract_text_generic(file_path)

	# Best-effort fallback for unknown text-like formats
	return _extract_text_generic(file_path)


def parse_project_documents(
	project_pdf_name: str,
	documents: dict,
	uploads_dir: Path,
) -> list[ParsedDocument]:
	file_candidates = [
		("final_brd", documents.get("finalBrdFile")),
		("mom", documents.get("momFile")),
		("pre_documents", documents.get("preDocumentsFile")),
		("transcripts", documents.get("transcriptsFile")),
		("project_pdf", project_pdf_name),
	]

	parsed_docs: list[ParsedDocument] = []
	seen_names: set[str] = set()

	for document_type, file_name in file_candidates:
		if not file_name:
			continue
		if file_name in seen_names:
			continue
		seen_names.add(file_name)

		file_path = uploads_dir / file_name
		if not file_path.exists():
			continue

		content = parse_file_content(file_path).strip()
		if not content:
			continue

		# Skip documents that are still mostly binary after cleaning
		if _looks_like_binary(content):
			continue

		parsed_docs.append(
			{
				"document_type": document_type,
				"file_name": file_name,
				"content": content,
			}
		)

	return parsed_docs